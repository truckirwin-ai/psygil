#!/usr/bin/env python3
"""Build 01_Project_Overview.docx for Psygil."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color palette ──────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x0D, 0x1B, 0x3E)   # deep navy
SLATE     = RGBColor(0x1E, 0x3A, 0x5F)   # section accent
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xF4, 0xF7, 0xFB)   # very light blue-grey for callout rows
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT    = RGBColor(0x2E, 0x86, 0xC1)   # mid-blue accent

OUTPUT = "/Users/truckirwin/Desktop/Projects/PSYCH_AI/docs/01_Project_Overview.docx"


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell background with a hex colour string (no #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def remove_cell_border(table):
    """Remove all inner/outer borders from a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:val'), 'none')
        tblBorders.append(node)
    tblPr.append(tblBorders)


def page_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)


def add_navy_heading(doc, text: str, level: int = 1):
    """Add a navy-background heading band."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    remove_cell_border(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, '0D1B3E')
    cell_margins(cell, top=120, bottom=120, left=160, right=160)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text.upper() if level == 1 else text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(13) if level == 1 else Pt(11)
    run.font.name = 'Calibri'
    doc.add_paragraph()  # breathing room


def add_subheading(doc, text: str):
    """Slate-coloured sub-section label."""
    table = doc.add_table(rows=1, cols=1)
    remove_cell_border(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, '1E3A5F')
    cell_margins(cell, top=80, bottom=80, left=160, right=160)
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    doc.add_paragraph()


def body(doc, text: str, space_after: bool = True):
    para = doc.add_paragraph(text)
    para.style = doc.styles['Normal']
    for run in para.runs:
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_TEXT
    if not space_after:
        para.paragraph_format.space_after = Pt(2)
    return para


def bullet(doc, text: str, bold_prefix: str = None):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        run_b = para.add_run(bold_prefix + '  ')
        run_b.bold = True
        run_b.font.size = Pt(10.5)
        run_b.font.name = 'Calibri'
        run_b.font.color.rgb = NAVY
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_TEXT


def spacer(doc, lines: int = 1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


# ── Cover block ────────────────────────────────────────────────────────────────

def add_cover(doc):
    table = doc.add_table(rows=1, cols=1)
    remove_cell_border(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, '0D1B3E')
    cell_margins(cell, top=320, bottom=320, left=320, right=320)

    # Product name
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run('PSYCHAI')
    r1.bold = True
    r1.font.size = Pt(32)
    r1.font.color.rgb = WHITE
    r1.font.name = 'Calibri'

    # Tagline
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('AI-Powered Evaluation Report Writing for Licensed Psychologists')
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0xAD, 0xC8, 0xE8)
    r2.font.name = 'Calibri'
    r2.italic = True

    # Meta line
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('Project Overview  |  Version 0.1  |  March 2026  |  CONFIDENTIAL')
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0x7F, 0x9F, 0xC0)
    r3.font.name = 'Calibri'

    spacer(doc, 2)


# ── Main document ──────────────────────────────────────────────────────────────

def build():
    doc = Document()
    page_margins(doc)

    # Default Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK_TEXT

    add_cover(doc)

    # ── 1. Executive Summary ─────────────────────────────────────────────────
    add_navy_heading(doc, '1. Executive Summary')
    body(doc,
        'Psygil is a secure, AI-assisted report writing platform built for licensed psychologists. '
        'It covers the complete evaluation workflow: from importing test results and source documents, '
        'through local PII removal, AI-assisted diagnostic analysis, structured review, and final '
        'publication as a polished Word document or PDF.'
    )
    body(doc,
        'Every patient record stays on the device. No identifiable data ever reaches the network. '
        'The platform is purpose-built for the real workflows of forensic and clinical psychologists '
        'who need speed, accuracy, and compliance in equal measure.'
    )
    spacer(doc)

    # ── 2. Problem Statement ─────────────────────────────────────────────────
    add_navy_heading(doc, '2. Problem Statement')
    body(doc,
        'Writing psychological evaluation reports is time-consuming and high-stakes. Psychologists '
        'typically spend 4 to 8 hours per report compiling scores, cross-referencing DSM-5-TR criteria, '
        'translating clinical notes into structured prose, and formatting the final document to meet '
        'referral or court standards.'
    )
    body(doc, 'The current process has three core problems:')
    bullet(doc, 'Manual assembly.  Test scores, interview notes, prior records, and Zoom transcripts live in separate files. Pulling them together is slow and error-prone.', bold_prefix=None)
    bullet(doc, 'Compliance risk.  Drafting in general-purpose word processors offers no guardrails around PII handling or HIPAA-adjacent best practices.')
    bullet(doc, 'Inconsistent voice.  Reports vary in tone, structure, and language depending on the day, the referral source, and how much time the psychologist had.')
    spacer(doc)
    body(doc,
        'Existing AI writing tools make the problem worse on the compliance front. They send raw '
        'clinical content to cloud servers before any scrubbing occurs. No licensed psychologist can '
        'accept that risk.'
    )
    spacer(doc)

    # ── 3. Solution Overview ─────────────────────────────────────────────────
    add_navy_heading(doc, '3. Solution Overview')
    body(doc,
        'Psygil solves all three problems without asking psychologists to change how they practice. '
        'It meets them in their existing workflow and handles the mechanical parts: data assembly, '
        'PII scrubbing, DSM-5-TR mapping, draft generation, and document formatting.'
    )
    body(doc, 'The architecture rests on three commitments:')
    bullet(doc, 'Local-first security.  PII removal runs on-device before any data leaves the machine. The AI sees anonymized clinical data only.')
    bullet(doc, 'Doctor in control.  The AI generates notes and drafts. The doctor reviews, modifies, approves, and signs off at every stage. Nothing publishes without explicit approval.')
    bullet(doc, 'Voice matching.  Doctors upload samples of their own prior reports. The system learns their preferred phrasing, structure, and style and writes to match it.')
    spacer(doc)

    # ── 4. Core Workflow ─────────────────────────────────────────────────────
    add_navy_heading(doc, '4. Core Workflow')
    body(doc,
        'Psygil structures every evaluation as a linear, gate-controlled workflow. Each stage '
        'requires an explicit doctor action before moving forward. No stage is skipped automatically.'
    )
    spacer(doc)

    stages = [
        ('Stage 1 -- Ingest',
         'The doctor imports source materials through a browser-based file tree. Accepted formats include '
         'Word documents, PDFs, plain text files, and Zoom meeting transcripts. Prior records and referral '
         'letters are added here as well.'),
        ('Stage 2 -- Cleanse',
         'Before any content reaches the AI, a local PII scrubbing engine strips patient-identifiable '
         'information: names, dates of birth, case numbers, addresses, and any other identifiers. This '
         'runs entirely on-device. The network never sees raw patient data.'),
        ('Stage 3 -- Analyze',
         'The AI receives the anonymized data and maps scores and observations to DSM-5-TR diagnostic '
         'criteria. It produces a structured set of preliminary diagnostic notes, flagging areas of '
         'ambiguity for the doctor to resolve.'),
        ('Stage 4 -- Review',
         'The doctor opens the diagnostic notes in a side-by-side panel alongside the source documents. '
         'They annotate, correct, or confirm each finding. This is the primary clinical judgment step. '
         'The AI does not modify anything without instruction.'),
        ('Stage 5 -- Draft',
         'The AI assembles the full evaluation report using the selected template, the doctor\'s writing '
         'style guide settings, and examples from the doctor\'s RAG library of prior reports. The output '
         'matches the doctor\'s preferred voice and structure.'),
        ('Stage 6 -- Edit',
         'The doctor edits the draft inline. They can type changes directly, issue voice commands '
         '("change the third paragraph to..."), or use the AI assistant for targeted rewrites. All '
         'changes appear in tracked-changes mode.'),
        ('Stage 7 -- Approve',
         'The doctor completes a final review pass. An approval gate requires explicit sign-off before '
         'the document advances to publication. The system logs the approval with a timestamp.'),
        ('Stage 8 -- Publish',
         'The finalized report exports as a formatted Word document or PDF. Output templates match '
         'common referral and court submission formats.'),
    ]

    for title, description in stages:
        add_subheading(doc, title)
        body(doc, description)
        spacer(doc)

    # ── 5. Key Features ──────────────────────────────────────────────────────
    add_navy_heading(doc, '5. Key Features')

    features = [
        ('Local PII Scrubbing Engine',
         'Runs entirely in-browser or on-device. Detects and removes names, dates, identifiers, and '
         'other PHI before any data leaves the machine. Configurable sensitivity levels per practice type.'),
        ('Browser-Based File System',
         'A visual file tree lets doctors navigate, import, and organize case files without leaving the app. '
         'Supports Word, PDF, and plain text formats.'),
        ('Word Doc Viewer, Editor, and Creator',
         'Full document creation and editing inside the platform. Doctors do not need to open a separate '
         'word processor at any point in the workflow.'),
        ('Template Library',
         'Comes with a set of standard evaluation report templates. Doctors import their own templates or '
         'customize existing ones. Templates control structure, headings, and required sections.'),
        ('Zoom Transcript Extraction',
         'Paste or import a Zoom transcript. The AI converts it into structured clinical notes formatted '
         'for inclusion in the evaluation report.'),
        ('RAG Writing Style System',
         'Doctors upload samples of their own prior reports. The retrieval-augmented generation system '
         'analyzes phrasing, sentence length, vocabulary, and structure. Drafts match the doctor\'s '
         'established voice.'),
        ('Writing Style Guide Settings',
         'Doctors set explicit rules: avoid em dashes, use active voice, avoid passive constructions, '
         'limit sentences to 20 words, and so on. The AI enforces these rules across every draft.'),
        ('DSM-5-TR Diagnostic Framework',
         'Test scores map automatically to DSM-5-TR criteria. The system flags which criteria are met, '
         'partially met, or not supported by the data. The doctor makes the final diagnostic determination.'),
        ('Voice Notes with AI Transcription',
         'Doctors record voice notes during or after a session. The AI transcribes and formats them. '
         'Inline voice commands let doctors direct edits without touching the keyboard.'),
        ('Side-by-Side Doc and Notes Panel',
         'The document viewer shows the report draft on one side and source materials or annotations '
         'on the other. Doctors never lose context while editing.'),
        ('Dark and Light Theme',
         'Full support for both display modes. Settings persist per user account.'),
        ('Word and PDF Export',
         'One-click export to formatted Word document or PDF. Templates control final layout, fonts, '
         'and header styles.'),
    ]

    for title, desc in features:
        add_subheading(doc, title)
        body(doc, desc)
        spacer(doc)

    # ── 6. Target Users ──────────────────────────────────────────────────────
    add_navy_heading(doc, '6. Target Users')
    body(doc,
        'Psygil is built for licensed psychologists who produce formal written evaluations. '
        'The primary users fall into two groups.'
    )
    spacer(doc)

    add_subheading(doc, 'Forensic Psychologists')
    body(doc,
        'Forensic psychologists produce evaluations for courts, attorneys, and corrections systems. '
        'Their reports follow strict structural requirements and face legal scrutiny. Speed matters. '
        'Consistency matters. Every word is on the record.'
    )
    spacer(doc)

    add_subheading(doc, 'Clinical Psychologists')
    body(doc,
        'Clinical psychologists evaluate patients for diagnosis, treatment planning, and insurance '
        'authorization. Their caseloads are high and administrative time cuts into clinical time. '
        'Psygil reduces the per-report burden without reducing clinical rigor.'
    )
    spacer(doc)

    body(doc,
        'Secondary users include neuropsychologists, school psychologists, and licensed psychological '
        'associates working under supervision. The platform accommodates multi-user practices with '
        'role-based access.'
    )
    spacer(doc)

    # ── 7. Guiding Principles ────────────────────────────────────────────────
    add_navy_heading(doc, '7. Guiding Principles')

    principles = [
        ('Privacy is not a feature, it is the foundation.',
         'Patient data never leaves the device in identifiable form. This is a hard requirement, '
         'not a setting. No exceptions are built into the architecture.'),
        ('The doctor decides. Always.',
         'AI generates, suggests, and drafts. The doctor confirms, modifies, and approves. '
         'The system does not advance a stage without explicit doctor action.'),
        ('Fit the real workflow.',
         'Psychologists already have a process. Psygil plugs into it. The platform does not '
         'require psychologists to learn a new methodology or reorganize their practice.'),
        ('Write like the doctor, not like a machine.',
         'The output should sound like the doctor\'s own work. The RAG system and style guides '
         'exist to eliminate generic AI prose and produce reports the doctor is proud to sign.'),
        ('Speed is clinical value.',
         'Every hour saved on report writing is an hour available for patients. The platform '
         'optimizes for time-to-publish without cutting corners on quality or compliance.'),
    ]

    for title, desc in principles:
        add_subheading(doc, title)
        body(doc, desc)
        spacer(doc)

    # ── Footer note ──────────────────────────────────────────────────────────
    spacer(doc)
    table = doc.add_table(rows=1, cols=1)
    remove_cell_border(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, '0D1B3E')
    cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'Psygil  |  Project Overview  |  v0.1  |  Confidential  |  '
        'Do not distribute outside core team'
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x7F, 0x9F, 0xC0)
    r.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Saved: {OUTPUT}')


if __name__ == '__main__':
    build()
