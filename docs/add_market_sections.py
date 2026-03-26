#!/usr/bin/env python3
"""Append WHAT EVERY COMPETITOR GETS WRONG and COMPETITIVE MOAT sections to 02_Market_Competitor_Research.docx."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import copy

DOC_PATH = "/Users/truckirwin/Desktop/Projects/PSYCH_AI/docs/02_Market_Competitor_Research.docx"

TEAL = RGBColor(0x17, 0xA5, 0x89)
GREY_LIGHT = RGBColor(0xCC, 0xCC, 0xCC)


def add_section_header(doc, text):
    """Teal 9pt bold uppercase section header matching existing doc style."""
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(18)
    sp.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = TEAL
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)


def add_body(doc, text):
    """11pt body paragraph."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, bold_prefix, body_text):
    """Bullet item with bold prefix and 11pt body text."""
    p = doc.add_paragraph(style='List Bullet')
    rb = p.add_run(bold_prefix + ': ')
    rb.bold = True
    rb.font.size = Pt(11)
    r = p.add_run(body_text)
    r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    return p


def build():
    doc = Document(DOC_PATH)

    # Save and remove last 3 paragraphs (footer line, divider, psychai.app)
    body = doc.element.body
    all_paras = body.findall(qn('w:p'))
    footer_elements = [copy.deepcopy(el) for el in all_paras[-3:]]
    for el in all_paras[-3:]:
        body.remove(el)

    # =========================================================================
    # SECTION: WHAT EVERY COMPETITOR GETS WRONG
    # =========================================================================
    add_section_header(doc, "WHAT EVERY COMPETITOR GETS WRONG")

    add_bullet(doc,
        "Note-takers, not report builders",
        "Every major competitor solves one problem: transcribe a session and format it as a SOAP note. "
        "A psychological evaluation report is 20 to 60 pages covering test score tables, "
        "criterion-by-criterion diagnostic analysis, differential diagnoses, and legal attestation. "
        "None of these tools touch that workflow. Psygil owns the entire process from raw test data "
        "to signed export."
    )

    add_bullet(doc,
        "Raw patient data goes to the cloud",
        "Every major competitor sends audio recordings or raw session text to their servers for "
        "processing. For forensic psychologists, that is a non-starter. Court cases, custody battles, "
        "criminal proceedings -- chain of custody matters and opposing counsel will probe it. "
        "Psygil's local-first PII engine means raw data never leaves the machine."
    )

    add_bullet(doc,
        "No voice matching",
        "PsychReport.AI claims to learn writing style from its user base, not from the individual "
        "doctor's own work. Psygil's RAG system indexes the doctor's own prior reports locally and "
        "writes the way that specific doctor writes. No cloud tool can do this without storing and "
        "processing the doctor's prior reports on their servers."
    )

    add_bullet(doc,
        "Manual test data entry",
        "Q-global, PAR, and Pearson export PDFs and CSVs. Every current tool makes the psychologist "
        "manually enter or paste scores. Psygil ingests those exports directly, normalizes scores, "
        "calculates confidence intervals, and flags validity concerns automatically."
    )

    add_bullet(doc,
        "Zoom transcripts ignored",
        "Telehealth is now standard. Forensic interviews, clinical interviews, and collateral calls "
        "happen on Zoom. Not one competitor has a workflow for ingesting those transcripts and "
        "converting them to structured clinical notes."
    )

    add_bullet(doc,
        "No legal defensibility",
        "No competitor has a Legal Reviewer agent. They produce documentation. Psygil produces "
        "documentation hardened against cross-examination -- every assertion sourced, every diagnosis "
        "criterion-cited, every causal claim flagged, with an immutable audit trail."
    )

    add_bullet(doc,
        "No workflow gates",
        "Every tool gives a draft and steps back. Psygil has four mandatory doctor approval gates "
        "creating a documented record that a licensed professional reviewed and approved each stage."
    )

    add_bullet(doc,
        "No style guide enforcement",
        "No tool enforces writing rules at generation time. Psygil's style guide engine means the "
        "doctor sets the rules once and the Writer Agent cannot violate them during generation."
    )

    add_bullet(doc,
        "Single-model architecture",
        "Every competitor runs one AI pass over the data. Psygil's eight-agent system means each "
        "step is handled by a specialist that cannot contaminate others. The Fact Checker cannot be "
        "overruled. The Legal Reviewer cannot be bypassed."
    )

    add_bullet(doc,
        "No forensic specialization",
        "The entire market optimizes for outpatient therapy notes. Forensic evaluation is a completely "
        "different discipline: multi-session, multi-instrument, court-admissible, often adversarial. "
        "Nobody has built for it."
    )

    add_body(doc,
        "The summary: every competitor solves the last ten minutes of a therapy session. "
        "Psygil solves the last ten hours of a forensic evaluation."
    )

    # =========================================================================
    # SECTION: COMPETITIVE MOAT
    # =========================================================================
    add_section_header(doc, "COMPETITIVE MOAT")

    add_body(doc,
        "Psygil's moat is the combination of local-first architecture and multi-agent legal hardening. "
        "Cloud-first competitors cannot adopt local PII processing without rebuilding their core "
        "infrastructure. Single-agent tools cannot add the legal defensibility layer without "
        "rearchitecting the entire pipeline. These are not features competitors can ship in a quarter."
    )

    add_body(doc,
        "The forensic psychology market is relationship-driven and reputation-sensitive. A tool that "
        "produces a report that fails in court destroys a practice. A tool that produces court-tested, "
        "audit-trailed, criterion-cited reports becomes indispensable. Switching costs are high once "
        "a doctor's prior reports are indexed as RAG examples and their style guide is configured."
    )

    add_body(doc,
        "The second moat is the audit trail itself. After two years of use, a doctor has a complete, "
        "timestamped record of every AI suggestion they accepted or rejected across hundreds of cases. "
        "That record is a defense exhibit. It is also a competitive lock-in that no competitor can replicate."
    )

    # Re-attach the saved footer elements
    for el in footer_elements:
        body.append(el)

    doc.save(DOC_PATH)
    print(f"Saved: {DOC_PATH}")


if __name__ == "__main__":
    build()
