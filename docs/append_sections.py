#!/usr/bin/env python3
"""Append Legal Hardening and Multi-Agent Architecture sections to 05_Architecture_Spec.docx."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = "/Users/truckirwin/Desktop/Projects/PSYCH_AI/docs/05_Architecture_Spec.docx"

# Colors from existing document
TEAL  = RGBColor(0x17, 0xA5, 0x89)   # 17A589 — teal footer/accent
NAVY  = RGBColor(0x0D, 0x3B, 0x5C)   # 0D3B5C — primary title navy
GREY_LIGHT = RGBColor(0xCC, 0xCC, 0xCC)  # CCCCCC — divider color
GREY_MID   = RGBColor(0x55, 0x55, 0x55)  # 555555 — subtitle


def add_divider(doc):
    """Light grey underscore divider line matching existing doc style."""
    p = doc.add_paragraph()
    run = p.add_run("_" * 72)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY_LIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)


def add_section_header(doc, label: str, title: str):
    """Teal 9pt bold uppercase label + navy 22pt bold title + divider."""
    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(20)
    sp.paragraph_format.space_after = Pt(0)

    # Teal uppercase label at 9pt bold
    p_label = doc.add_paragraph()
    r = p_label.add_run(label.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = TEAL
    p_label.paragraph_format.space_before = Pt(0)
    p_label.paragraph_format.space_after = Pt(2)

    # Navy 22pt bold title
    p_title = doc.add_paragraph()
    r2 = p_title.add_run(title)
    r2.bold = True
    r2.font.size = Pt(22)
    r2.font.color.rgb = NAVY
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(6)

    add_divider(doc)


def add_subsection_heading(doc, text: str):
    """Bold 11pt subsection heading, no explicit color (inherits)."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)


def add_body(doc, text: str):
    """11pt body paragraph."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_bullet(doc, text: str, bold_prefix: str = None):
    """Bullet list item at 11pt."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        rb = p.add_run(bold_prefix + ": ")
        rb.bold = True
        rb.font.size = Pt(11)
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    return p


def build():
    doc = Document(DOC_PATH)

    # =========================================================================
    # SECTION A: LEGAL HARDENING ARCHITECTURE
    # =========================================================================
    add_section_header(
        doc,
        label="Section A",
        title="Legal Hardening Architecture"
    )

    add_body(doc,
        "Psygil reports are submitted to courts, licensing boards, and legal proceedings. "
        "Every sentence is subject to cross-examination. The system is designed from the ground up "
        "to produce reports that are defensible under adversarial scrutiny. Legal defensibility is not "
        "a review step added at the end. It is a property of the entire pipeline."
    )

    # --- 1. Source Traceability ---
    add_subsection_heading(doc, "1. Source Traceability")
    add_body(doc,
        "Every clinical assertion in a final report links to a specific source data point. "
        "No statement enters the report without a traceable origin in the underlying data."
    )
    add_bullet(doc,
        "The system maintains a citation graph mapping each claim to its source data point, "
        "to the test or instrument that produced it, and to the administrator notes recorded at the time."
    )
    add_bullet(doc,
        "A statement such as 'the patient demonstrates impaired executive function' requires "
        "a linked T-score or documented behavioral observation. The link is stored in the report "
        "metadata and is auditable on demand."
    )
    add_bullet(doc,
        "The Writer Agent cannot generate unsupported assertions. When the data packet contains "
        "no source that supports a claim, the agent flags the gap and holds the sentence for "
        "doctor resolution. It does not invent supporting evidence."
    )

    # --- 2. Diagnostic Citation Requirements ---
    add_subsection_heading(doc, "2. Diagnostic Citation Requirements")
    add_body(doc,
        "Every diagnosis in the report references the specific DSM-5-TR criteria met, cited inline "
        "with the criterion letter and number. Vague diagnostic language that cannot be traced to "
        "explicit criteria is blocked."
    )
    add_bullet(doc,
        "Criterion citations appear inline within the diagnostic section, for example: "
        "'Criterion A1, A3, and B2 of DSM-5-TR 296.89 are met based on the following data.'"
    )
    add_bullet(doc,
        "Differential diagnoses are documented with explicit written rationale for ruling each "
        "condition in or out. The rationale references the data, not the clinician's general impression."
    )
    add_bullet(doc,
        "The system flags when the available data supports a diagnosis, a rule-out, or only a "
        "provisional diagnosis. These three categories carry distinct language requirements "
        "and different levels of certainty."
    )
    add_bullet(doc,
        "ICD-10-CM codes are required fields. The export pipeline blocks finalization if any "
        "diagnosis is missing a corresponding ICD-10-CM code."
    )

    # --- 3. Language Defensibility Engine ---
    add_subsection_heading(doc, "3. Language Defensibility Engine")
    add_body(doc,
        "The Legal Reviewer Agent reads every sentence of the draft report as a hostile attorney would. "
        "It identifies language patterns that create legal exposure and flags each one with a specific "
        "remediation option."
    )
    add_body(doc, "The following language categories are subject to automatic flagging:")
    add_bullet(doc,
        "Hedge language that implies uncertainty without clinical justification -- phrases such as "
        "'appears to,' 'may suggest,' and 'seems consistent with' -- is flagged unless the uncertainty "
        "is explicitly documented in the source data and the hedging is intentional.",
        bold_prefix="Unsupported hedging"
    )
    add_bullet(doc,
        "Absolute language used without sufficient evidentiary support -- 'definitively,' 'certainly,' "
        "'conclusively' -- is flagged for review.",
        bold_prefix="Overreach language"
    )
    add_bullet(doc,
        "Causal claims that lack supporting evidence in the data record, such as 'the abuse caused "
        "the symptom presentation,' are flagged as unsupported causal assertions.",
        bold_prefix="Unsupported causation"
    )
    add_bullet(doc,
        "Speculative future predictions, such as 'will likely reoffend,' are flagged unless the "
        "statement is directly supported by a validated actuarial risk instrument with its score "
        "and normative sample cited.",
        bold_prefix="Actuarial predictions"
    )
    add_body(doc,
        "Each flag produces a category label, the specific sentence or phrase at issue, and a "
        "suggested remediation. The agent does not rewrite the text. The doctor resolves each flag."
    )

    # --- 4. Immutable Audit Trail ---
    add_subsection_heading(doc, "4. Immutable Audit Trail")
    add_body(doc,
        "The audit trail is the system's primary defense against the claim that the AI, rather than "
        "the licensed clinician, authored the report. Every action taken on every document is logged "
        "in an append-only record."
    )
    add_bullet(doc,
        "Every document ingested into the system receives a SHA-256 hash recorded at intake. "
        "Chain of custody is logged from first import through final export."
    )
    add_bullet(doc,
        "Every AI suggestion is logged with a timestamp, the agent ID, the model version used, "
        "and a hash of the prompt that produced the output."
    )
    add_bullet(doc,
        "Every human edit is logged with a timestamp, the user ID of the clinician, and a record "
        "of what changed."
    )
    add_bullet(doc,
        "Every approval gate is logged with the doctor's ID, a timestamp, and the full attestation "
        "text accepted at that gate."
    )
    add_bullet(doc,
        "The audit log is append-only and tamper-evident. Each log entry includes a hash of the "
        "preceding entry, forming a chain that detects any retroactive modification."
    )
    add_bullet(doc,
        "On export, a separate audit report is available that displays the complete chronological "
        "history of the document from intake to finalization."
    )

    # --- 5. Version Control and Diff Tracking ---
    add_subsection_heading(doc, "5. Version Control and Diff Tracking")
    add_body(doc,
        "Every save operation creates a version snapshot. No prior state of the document is "
        "overwritten. Diffs between versions are stored and viewable within the application."
    )
    add_bullet(doc, "The doctor can restore any prior version at any point before export.")
    add_bullet(doc,
        "On export, the version number and finalization date are embedded in the document metadata."
    )
    add_bullet(doc,
        "If the report is modified after a prior export, the version number increments and the "
        "change is logged in the audit trail."
    )

    # --- 6. Export Legal Standards ---
    add_subsection_heading(doc, "6. Export Legal Standards")
    add_body(doc,
        "Reports leave the system only when they meet a defined set of legal and administrative "
        "standards. The export pipeline enforces these requirements automatically."
    )
    add_bullet(doc,
        "Exported reports carry embedded metadata: author name, export date, document version, "
        "and Psygil version number."
    )
    add_bullet(doc,
        "A digital signature block with the following attestation language is required at the "
        "bottom of every exported report: 'I attest that I have reviewed, edited, and approve "
        "the contents of this report. AI assistance was used in drafting under my direct supervision.'"
    )
    add_bullet(doc,
        "Export requires explicit doctor sign-off at the final approval gate. The system does not "
        "produce a finalized export without this confirmation."
    )
    add_bullet(doc,
        "Unapproved drafts carry a watermark: 'DRAFT - NOT FOR DISTRIBUTION.' The watermark is "
        "removed only after the approval gate is cleared."
    )

    # =========================================================================
    # SECTION B: MULTI-AGENT SYSTEM ARCHITECTURE
    # =========================================================================
    add_section_header(
        doc,
        label="Section B",
        title="Multi-Agent System Architecture"
    )

    add_body(doc,
        "Each major step in the evaluation workflow is handled by a dedicated AI agent. Agents do "
        "not share state directly. They communicate only through the Orchestrator. No agent can "
        "override another agent's output. Each agent has a defined identity, a bounded mandate, "
        "and hard guardrails that cannot be overridden by any other agent or by the Orchestrator."
    )
    add_body(doc,
        "Doctor approval gates are built into the workflow at four checkpoints. If any agent "
        "raises a hard stop flag, the workflow pauses and the doctor is notified before the "
        "process resumes."
    )

    # --- Agent 1: The Orchestrator ---
    add_subsection_heading(doc, "Agent 1: The Orchestrator")

    p = doc.add_paragraph()
    rb = p.add_run("Soul: ")
    rb.bold = True
    rb.font.size = Pt(11)
    r = p.add_run(
        "Workflow coordinator. Moves the evaluation through stages. Has no clinical opinions."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Mandate: ")
    rb.bold = True
    rb.font.size = Pt(11)
    r = p.add_run(
        "Manages stage transitions, routes work to domain agents, enforces approval gates, "
        "and maintains the master job state for the active evaluation."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Guardrails: ")
    rb.bold = True
    rb.font.size = Pt(11)
    r = p.add_run(
        "Cannot generate clinical content. Cannot approve or reject diagnostic conclusions. "
        "Cannot modify report text. Can only advance, pause, or return a stage."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Inputs: ")
    rb.bold = True
    rb.font.size = Pt(11)
    r = p.add_run("Stage completion signals from domain agents, doctor approvals.")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Outputs: ")
    rb.bold = True
    rb.font.size = Pt(11)
    r = p.add_run("Stage transitions, status updates, alerts.")
    r.font.size = Pt(11)

    # --- Agent 2: The Ingestor ---
    add_subsection_heading(doc, "Agent 2: The Ingestor")

    p = doc.add_paragraph()
    rb = p.add_run("Soul: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run("Data intake specialist. Precise, methodical, no interpretation.")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Mandate: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Accepts raw documents including test results, records, referral letters, and Zoom "
        "transcripts. Runs PII detection. Normalizes and structures data. Outputs a clean, "
        "de-identified data packet."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Guardrails: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Cannot interpret data. Cannot make clinical inferences. Cannot generate prose. "
        "If PII detection confidence falls below the configured threshold, the agent flags the "
        "document for human review before proceeding. Must reject documents it cannot parse "
        "rather than guess at their contents."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Inputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run("Raw files from doctor.")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Outputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run("Structured, de-identified data packet with field-level provenance.")
    r.font.size = Pt(11)

    # --- Agent 3: The Test Data Analyst ---
    add_subsection_heading(doc, "Agent 3: The Test Data Analyst")

    p = doc.add_paragraph()
    rb = p.add_run("Soul: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Psychometrician. Speaks in scores, percentiles, and confidence intervals. "
        "No clinical speculation."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Mandate: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Receives the structured data packet. Interprets all psychometric scores: converts raw "
        "scores, applies norms, calculates confidence intervals, identifies score patterns, and "
        "flags outliers and validity concerns such as elevated MMPI validity scales suggesting "
        "response distortion."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Guardrails: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Cannot make DSM diagnoses. Cannot recommend treatment. Cannot generate narrative prose "
        "beyond technical score descriptions. Must cite the specific normative sample used for "
        "every interpretation. Must flag any score that falls outside the validated age or "
        "demographic range for its normative sample."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Inputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run("Structured data packet from Ingestor.")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Outputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Psychometric summary table, score interpretation notes, and validity flags."
    )
    r.font.size = Pt(11)

    # --- Agent 4: The Diagnostician ---
    add_subsection_heading(doc, "Agent 4: The Diagnostician")

    p = doc.add_paragraph()
    rb = p.add_run("Soul: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Careful, conservative clinician. Does not overreach. Never diagnoses on insufficient data."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Mandate: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Maps test analyst output and clinical observations to DSM-5-TR criteria. Formulates "
        "preliminary diagnostic impressions with explicit criterion-by-criterion analysis. "
        "Produces a differential diagnosis with documented rationale for each condition "
        "included or excluded."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Guardrails: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Cannot generate report prose. Cannot override a doctor's diagnosis. Must cite specific "
        "DSM-5-TR criteria for every diagnostic impression. Must document criteria that are NOT "
        "met as well as those that are. Must flag any diagnosis unsupported by two or more "
        "independent data sources. Cannot diagnose from a single test result alone."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Inputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run("Psychometric summary, clinical observations, and interview notes.")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Outputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Diagnostic impression memo with full DSM-5-TR criterion mapping and differential "
        "diagnosis with rationale."
    )
    r.font.size = Pt(11)

    # --- Agent 5: The Writer ---
    add_subsection_heading(doc, "Agent 5: The Writer")

    p = doc.add_paragraph()
    rb = p.add_run("Soul: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Precise, clear clinical writer. Serves the doctor's voice, not its own."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Mandate: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Assembles the evaluation report from structured inputs. Uses the doctor's RAG examples "
        "for voice matching. Enforces writing style guide rules. Generates section by section "
        "with streaming output. Every assertion links to source data in the citation graph."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Guardrails: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Cannot introduce clinical content not present in the inputs. Cannot modify diagnostic "
        "conclusions. Cannot soften or amplify findings beyond what the data supports. Must flag "
        "any sentence it generates that lacks a traceable source citation. Cannot trigger export. "
        "Produces drafts only."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Inputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Diagnostic memo, psychometric summary, clinical notes, doctor's RAG examples, "
        "and writing style guide."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Outputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run("Draft report sections with inline source citations.")
    r.font.size = Pt(11)

    # --- Agent 6: The Legal Reviewer ---
    add_subsection_heading(doc, "Agent 6: The Legal Reviewer")

    p = doc.add_paragraph()
    rb = p.add_run("Soul: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run("Adversarial reader. Reads every sentence as a hostile lawyer would.")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Mandate: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Reviews the draft report for legally vulnerable language. Flags unsupported assertions, "
        "problematic hedge language, causal claims without basis, and speculative predictions. "
        "Produces a redline report with specific flagged items and remediation options."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Guardrails: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Cannot rewrite text. Can only flag and suggest. Cannot make clinical judgments. "
        "Cannot approve the report. The agent errs on the side of flagging. False positives "
        "are acceptable. False negatives are not."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Inputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run("Draft report from Writer.")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Outputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Redline report with flagged items, flag categories, and suggested remediations."
    )
    r.font.size = Pt(11)

    # --- Agent 7: The Editor ---
    add_subsection_heading(doc, "Agent 7: The Editor")

    p = doc.add_paragraph()
    rb = p.add_run("Soul: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run("Copy editor with clinical literacy. Serves clarity and the doctor's voice.")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Mandate: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Applies writing style guide rules. Checks consistency, eliminates redundancy, and "
        "improves sentence structure. Does not alter clinical content -- only improves how "
        "it reads."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Guardrails: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Cannot change clinical conclusions. Cannot add or remove diagnostic content. Cannot "
        "override Legal Reviewer flags. Must preserve all source citations. All changes are "
        "tracked, not applied silently."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Inputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run("Draft report after Legal Reviewer pass.")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Outputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run("Edited draft with tracked changes.")
    r.font.size = Pt(11)

    # --- Agent 8: The Fact Checker ---
    add_subsection_heading(doc, "Agent 8: The Fact Checker")

    p = doc.add_paragraph()
    rb = p.add_run("Soul: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run("Skeptic. Trusts nothing it cannot verify.")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Mandate: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Verifies all factual claims in the report against the source data packet. Checks that "
        "every score cited matches the input data. Checks that every DSM-5-TR criterion citation "
        "is accurate. Checks that demographic details are internally consistent. Checks that "
        "dates are correct throughout the document."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Guardrails: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Cannot make editorial judgments. Cannot change text. Can only flag discrepancies. "
        "A discrepancy between the report and source data is a hard stop. The report cannot "
        "advance to final review until every discrepancy is resolved."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Inputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Final draft, original source data packet, and DSM-5-TR reference database."
    )
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    rb = p.add_run("Outputs: ")
    rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(
        "Verification report with pass/fail status per section and a list of flagged discrepancies."
    )
    r.font.size = Pt(11)

    # --- Agent Interaction Rules ---
    add_subsection_heading(doc, "Agent Interaction Rules")
    add_body(doc,
        "The following rules govern all agent behavior and communication across the system:"
    )
    add_bullet(doc,
        "Agents communicate through the Orchestrator only. No direct agent-to-agent calls are permitted."
    )
    add_bullet(doc,
        "Each agent receives only the data it needs for its specific task. "
        "Least-privilege data access is enforced at the Orchestrator level."
    )
    add_bullet(doc,
        "Each agent logs its full reasoning chain, not just its final output. "
        "The reasoning log is part of the audit trail."
    )
    add_bullet(doc,
        "No agent can modify another agent's output directly. An agent can flag another "
        "agent's output for review. Only the doctor resolves flags."
    )
    add_bullet(doc,
        "Doctor approval is required at four points: after the Ingestor completes, after the "
        "Diagnostician completes, after the Legal Reviewer completes, and before export."
    )
    add_bullet(doc,
        "If any agent raises a hard stop flag, the workflow pauses immediately and the doctor "
        "is notified. The workflow does not resume until the doctor takes explicit action."
    )

    # Final divider
    add_divider(doc)

    doc.save(DOC_PATH)
    print(f"Saved: {DOC_PATH}")


if __name__ == "__main__":
    build()
